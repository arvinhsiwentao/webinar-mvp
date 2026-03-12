"""
Generate comprehensive tracking/analytics audit report as DOCX
Rewritten for non-technical marketing/ops readers
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Global styles ──────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Adjust heading sizes
doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)

GOLD = RGBColor(0xB8, 0x95, 0x3F)
RED = RGBColor(0xCC, 0x00, 0x00)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
GRAY = RGBColor(0x66, 0x66, 0x66)

# ── Helpers ────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_callout(doc, text, color=GOLD):
    """Add a highlighted callout box (simulated with indented bold paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    run.italic = True
    return p

def add_bold_paragraph(doc, bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_code_block(doc, code_text, caption=''):
    """Add a code snippet with gray background feel."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


# ╔══════════════════════════════════════════════════════════╗
# ║                      封面                                ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Webinar 直播平台')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = GOLD

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('数据追踪 · 埋点审计报告')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    '本报告用白话文写成，非技术人员也能完全理解\n\n'
    '适用对象：行销投手 · 运营经理 · 数据分析师 · 老板\n'
    '用途：核对 GTM / GA4 / Google Ads 的追踪设置\n\n'
    '报告日期：2026-03-12'
)
run.font.size = Pt(12)
run.font.color.rgb = GRAY

doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║                     术语表                               ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('先看这里：关键术语速查表', level=1)

doc.add_paragraph(
    '这份报告会用到一些数据追踪领域的专有名词。'
    '不用死记硬背，遇到不懂的随时翻回来查：'
)

add_table(doc,
    ['术语', '白话翻译', '为什么重要'],
    [
        ['埋点 / 事件追踪',
         '在网页的关键位置放上「记录器」，\n当用户做某个动作时自动记一笔',
         '没有埋点 = 瞎投广告，\n不知道哪个广告带来了购买'],
        ['GA4\n(Google Analytics 4)',
         'Google 的免费数据分析工具，\n帮你看网站的访客行为',
         '你在 GA4 后台看到的所有数据，\n都来自网页上的埋点'],
        ['GTM\n(Google Tag Manager)',
         'Google 的「标签管家」，\n帮你统一管理网页上的各种追踪代码',
         '我们的埋点数据先发给 GTM，\nGTM 再转发给 GA4 和 Google Ads'],
        ['转化事件',
         '你最在意的用户动作，\n例如「注册」「购买」',
         'Google Ads 靠转化事件来判断\n哪些广告有效、该把预算花在哪'],
        ['gclid',
         'Google 广告点击 ID，\n用户点广告时 Google 自动带上',
         '有了它，才能知道「这笔购买\n是从哪个广告关键字来的」'],
        ['UTM 参数\n(utm_source 等)',
         '你手动加在网址后面的追踪标记，\n例如 ?utm_source=facebook',
         '帮你区分流量来源：\nGoogle 广告、Facebook、EDM 等'],
        ['Enhanced Conversions\n(增强型转化)',
         '把用户的 email/手机传给 Google，\n让 Google 更准确地对比\n「谁看了广告→谁最终购买」',
         'Apple 隐私政策越来越严，\n没有这个功能转化数据会不准'],
        ['ROAS',
         '广告投入产出比\n= 广告带来的收入 ÷ 广告花费',
         '例如 ROAS 5 = 花 1 元广告费\n赚回 5 元，是投放的核心指标'],
        ['再营销 / Retargeting',
         '对「来过但没买」的人再投广告',
         '比找新客人便宜很多，\n转化率通常高 3~5 倍'],
        ['dataLayer',
         '网页上的一个「数据信箱」，\n埋点数据先放进去，GTM 来取',
         '你在 GTM 里设 Tag 时，\n数据就是从这里读的'],
        ['Measurement Protocol',
         '从服务器直接发数据给 GA4\n（不经过用户的浏览器）',
         '用户关掉浏览器也不怕漏记，\n是最可靠的追踪方式'],
    ]
)

doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║                     目录                                 ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('目录', level=1)
toc_items = [
    ('一', '这份报告在说什么？'),
    ('二', '追踪是怎么运作的（一张图看懂）'),
    ('三', '用户旅程：每一步我们记了什么'),
    ('', '   ① 着陆页（用户第一眼看到的页面）'),
    ('', '   ② 注册'),
    ('', '   ③ 候场等待'),
    ('', '   ④ 观看直播'),
    ('', '   ⑤ 直播结束 / 销售页'),
    ('', '   ⑥ 付款 / 购买完成'),
    ('四', '转化事件总表：Google Ads 最在意的 6 个动作'),
    ('五', '广告来源追踪：我们是怎么记住「用户从哪来」的'),
    ('六', '再营销受众：可以圈出哪些人群再投广告'),
    ('七', '健康检查：目前做得好的 & 还缺什么'),
    ('', '   P0 — 必须马上补（没有会亏钱）'),
    ('', '   P1 — 建议尽快补（能明显提升效果）'),
    ('', '   P2 — 锦上添花（有更好）'),
    ('八', 'GTM / GA4 配置端要做的事（给技术同事的清单）'),
]
for num, text in toc_items:
    p = doc.add_paragraph()
    if num:
        run = p.add_run(f'{num}. ')
        run.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║              一、这份报告在说什么                          ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('一、这份报告在说什么？', level=1)

doc.add_paragraph(
    '想象一下：你花了 1 万块投 Google 广告，有 100 个人点进来看了直播，其中 5 个人付费购买了课程。'
)
doc.add_paragraph(
    '但如果网站上没有做好「数据追踪」，你就无法知道：'
)
bullets = [
    '这 5 个付费客户，是从哪个广告关键字点进来的？',
    '剩下 95 个没买的人，卡在哪一步走掉了？是没看完直播？还是看完了但没点购买？',
    '下次投广告，应该加大哪个关键字的预算？',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    '「埋点」就是在网页的关键位置放上看不见的记录器。'
    '每当用户做了一个重要动作（例如注册、点击购买、看完直播），记录器就自动通知 Google Analytics，让我们事后能分析。'
)

add_callout(doc,
    '这份报告完整盘点了网站上的 19 个追踪事件，涵盖用户从看到广告到最终付款的每一步。'
    '同时标注了哪些追踪已经做好、哪些还缺，并按优先级排出修复建议。'
)

doc.add_paragraph(
    '报告覆盖范围：\n'
    '  • 19 个独立追踪事件（分布在 22 个触发点）\n'
    '  • 7 个用户旅程阶段：看到广告 → 着陆页 → 注册 → 候场 → 直播 → 结束页 → 付款\n'
    '  • 6 个关键转化事件（Google Ads 靠这些来优化出价）\n'
    '  • 7 种再营销受众（可以圈出来再投广告的人群）'
)


# ╔══════════════════════════════════════════════════════════╗
# ║              二、追踪是怎么运作的                          ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('二、追踪是怎么运作的（一张图看懂）', level=1)

doc.add_paragraph('整个追踪的运作方式可以用一条「数据流水线」来理解：')

doc.add_paragraph(
    '用户在网页上做动作（例如点击「立即注册」）\n'
    '        ↓\n'
    '网页代码自动把这个动作记录到一个「数据信箱」（dataLayer）\n'
    '        ↓\n'
    'Google Tag Manager（GTM）从信箱取出数据\n'
    '        ↓\n'
    'GTM 把数据分别送给 GA4（看报表）和 Google Ads（优化广告）'
)

add_callout(doc,
    '你不需要理解代码细节。只要知道：网页上每发生一个重要事件，都会自动通知 Google。'
    '你在 GA4 后台看到的数据、Google Ads 自动优化出价用的转化数据，都来自这条流水线。'
)

doc.add_heading('广告来源是怎么记住的？', level=2)
doc.add_paragraph(
    '当用户通过 Google 广告点进来时，网址后面会自动带上一串代码（叫做 gclid）。\n\n'
    '我们的网站会在用户进来的第一秒就把这串代码存下来（存两份：浏览器临时存储 + Cookie）。\n\n'
    '之后不管用户在网站上做什么——注册、看直播、付款——每一个重要事件都会自动带上这串代码，'
    '这样 Google Ads 就能把「这笔购买」对应回「那个广告点击」。\n\n'
    'UTM 参数（utm_source、utm_campaign 等）也是同样的原理，用来追踪非 Google 的流量来源，'
    '例如 Facebook、EDM 邮件等。'
)


doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║       三、用户旅程：每一步我们记了什么                       ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('三、用户旅程：每一步我们记了什么', level=1)

doc.add_paragraph(
    '下面按照用户从看到广告到最终付款的顺序，逐步说明每一步追踪了什么。\n'
    '表格中的「事件名」就是会出现在你 GA4 后台的名字，直接用来核对。'
)

# ── 总览表 ──
doc.add_heading('总览：用户旅程追踪地图', level=2)

add_table(doc,
    ['步骤', '用户在做什么', '我们记录了什么', '对广告的意义'],
    [
        ['① 着陆页',
         '用户点了广告，第一次看到我们的页面',
         '看了多少内容（滚动深度）\n有没有点「报名」按钮',
         '衡量着陆页吸引力\n判断广告和页面是否匹配'],
        ['② 注册',
         '用户填表注册',
         '注册成功',
         '最重要的转化之一\nGoogle Ads 用这个来找更多类似的人'],
        ['③ 候场等待',
         '用户到了等待页，\n等直播开始',
         '有没有来等待\n等了多久、有没有中途离开\n最终有没有进入直播间',
         '分析注册后的出席率\n找出等待页的流失原因'],
        ['④ 观看直播',
         '用户在看直播',
         '每 60 秒记录一次（还在看吗？）\n看到了百分之几\n有没有发弹幕\n有没有看到购买弹窗\n有没有点「购买」',
         '最细致的互动数据\n判断直播内容哪里最吸引人\n哪个时间点最容易成交'],
        ['⑤ 结束页',
         '直播结束，用户看到销售页',
         '有没有看完整场直播\n有没有点购买按钮\n有没有分享给朋友',
         '最后的转化推动力\n衡量直播结束后的销售效果'],
        ['⑥ 付款完成',
         '用户完成 Stripe 付款',
         '交易金额、交易编号\n订单状态确认（双重保险）',
         '最核心的数据\n用来算广告投入产出比（ROAS）'],
    ]
)

doc.add_page_break()

# ── helper: event detail card ──
def add_event_card(doc, event_name, trigger, value, params, code, is_conversion=False):
    """Add a full event detail card with parameter table."""
    # Event heading
    conv_tag = '  [转化事件 ✓]' if is_conversion else ''
    doc.add_heading(f'事件：{event_name}{conv_tag}', level=3)
    # Trigger + value
    add_bold_paragraph(doc, '什么时候触发：', trigger)
    add_bold_paragraph(doc, '运营价值：', value)
    # Param table
    p = doc.add_paragraph()
    run = p.add_run('参数明细（在 GA4 后台 → 事件 → 参数 中核对）：')
    run.bold = True
    run.font.size = Pt(10)
    add_table(doc,
        ['参数名\n（GA4 显示名称）', '白话含义', '值的格式 / 范例'],
        params,
    )
    # Code
    add_code_block(doc, code, '代码对照：')


# ── ① 着陆页 ──
doc.add_heading('① 着陆页（用户第一眼看到的页面）', level=2)
doc.add_paragraph(
    '这是用户点了广告后看到的第一个页面。我们想知道：用户有没有认真看？有没有被吸引？\n'
    '文件位置：src/app/(public)/page.tsx'
)

add_event_card(doc,
    event_name='c_scroll_depth',
    trigger='用户往下滚动页面时，每滚过 10% 记一次（10%, 20%, 30%... 直到 100%），每个里程碑只记一次',
    value='如果大多数人只看了 20% 就走了，说明页面开头不够吸引人；如果很多人到了 80% 以上，说明内容抓人',
    params=[
        ['percent', '用户滚动到了页面的百分之几', '数字：10, 20, 30 ... 100'],
        ['page', '哪个页面（目前固定是着陆页）', '固定值："landing"'],
    ],
    code="trackGA4('c_scroll_depth', { percent: 50, page: 'landing' });\n"
         "// 意思：用户看了着陆页 50% 的内容",
)

add_event_card(doc,
    event_name='c_signup_button_click',
    trigger='用户点了「观看讲座」报名按钮（页面顶部 Hero 区域或页面底部都有一个按钮）',
    value='知道哪个位置的按钮更多人点，可以优化页面布局；也是注册漏斗的第一步',
    params=[
        ['button_position', '用户点的是哪个按钮', '"hero" = 页面顶部大按钮\n"footer" = 页面底部按钮'],
        ['webinar_id', '对应的直播场次编号', '固定值："1"（目前只有一场）'],
    ],
    code="trackGA4('c_signup_button_click', {\n"
         "  button_position: 'hero',\n"
         "  webinar_id: '1'\n"
         "});",
)


# ── ② 注册 ──
doc.add_heading('② 注册', level=2)
doc.add_paragraph(
    '用户填写表单并成功注册后，会触发一个「sign_up」事件。'
    '这是 Google 官方推荐的标准事件名，GA4 会自动识别。\n'
    '文件位置：src/components/registration/useRegistrationForm.ts'
)

add_event_card(doc,
    event_name='sign_up',
    trigger='注册表单提交成功（服务器确认没问题后才触发，不是一点提交就算）',
    value='最重要的转化事件之一。Google Ads 会用这个事件来优化「找到更多会注册的人」',
    params=[
        ['method', '注册来源，标明用户从哪个按钮点进来注册的',
         '"webinar_registration_hero"\n"webinar_registration_footer"\n"webinar_registration"（来源不明时）'],
        ['webinar_id', '直播场次编号', '"1"'],
        ['--- 以下为自动附加的归因参数 ---', '', ''],
        ['gclid', 'Google 广告点击 ID（如果有的话）', '"EAIaIQob..." 或 空'],
        ['utm_source', '流量来源', '"google" / "facebook" 等'],
        ['utm_medium', '流量媒介', '"cpc" / "email" 等'],
        ['utm_campaign', '广告活动名称', '"spring_sale" 等'],
        ['utm_content', '广告素材标记', '"video_ad_a" 等'],
    ],
    code="trackGA4('sign_up', {\n"
         "  method: 'webinar_registration_hero',\n"
         "  webinar_id: '1'\n"
         "  // + gclid, utm_source, utm_medium, utm_campaign, utm_content (自动附加)\n"
         "});",
    is_conversion=True,
)

add_callout(doc,
    '注意：目前 sign_up 事件只传了 method 和 webinar_id，没有传用户的 email 和手机。'
    '这会影响 Google Ads 的 Enhanced Conversions 功能（后面第七章 P0-1 会详细说明）。',
    RED
)


# ── ③ 候场等待 ──
doc.add_heading('③ 候场等待', level=2)
doc.add_paragraph(
    '注册后，用户会到一个等待页面，等待直播开始。'
    '这个阶段的追踪非常完善，可以回答很多运营问题。\n'
    '文件位置：src/app/(public)/webinar/[id]/lobby/page.tsx'
)

add_event_card(doc,
    event_name='c_lobby_entered',
    trigger='用户打开了候场等待页（页面加载完成后触发，每次访问只触发一次）',
    value='用来算「出席率」：注册了 100 人，但只有 60 人打开等待页 → 出席率 60%',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['webinar_state', '当前直播状态\n（用户到达候场页时，直播处于什么阶段）',
         '"pre_live" = 还没开始\n"live" = 正在直播\n"post_live" = 已结束\n"standard" = 普通场次'],
    ],
    code="trackGA4('c_lobby_entered', {\n"
         "  webinar_id: '1',\n"
         "  webinar_state: 'pre_live'\n"
         "});",
)

add_event_card(doc,
    event_name='c_enter_live',
    trigger='用户从候场页成功进入直播间（有 4 种进入方式，都会记录）',
    value='衡量候场页的转化率：来等待的 60 人里，最终有多少人真的进入直播间？',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['entry_method', '用户是怎么进入直播间的',
         '"redirect_live" = 打开候场页时直播已开始，自动跳转\n'
         '"countdown_auto" = 倒计时归零，自动跳转\n'
         '"button" = 用户手动点了「进入直播间」按钮'],
        ['--- 以下为自动附加的归因参数 ---', '', ''],
        ['gclid', 'Google 广告点击 ID', '同 sign_up'],
        ['utm_source / medium / campaign / content', '广告来源参数', '同 sign_up'],
    ],
    code="trackGA4('c_enter_live', {\n"
         "  webinar_id: '1',\n"
         "  entry_method: 'countdown_auto'\n"
         "  // + gclid, utm_* (自动附加)\n"
         "});",
    is_conversion=True,
)

add_event_card(doc,
    event_name='c_lobby_duration',
    trigger='用户离开候场页时触发（不管是进入直播间还是中途放弃都会触发）',
    value='分析用户等待耐心：平均等了多久？等太久的人是不是更容易流失？',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['duration_sec', '用户在候场页总共待了几秒', '数字，例如：300（= 5分钟）'],
        ['exit_type', '用户最终是进了直播间还是放弃了',
         '"enter_live" = 成功进入直播间\n"abandon" = 中途离开了'],
    ],
    code="trackGA4('c_lobby_duration', {\n"
         "  webinar_id: '1',\n"
         "  duration_sec: 300,\n"
         "  exit_type: 'enter_live'\n"
         "});",
)

add_event_card(doc,
    event_name='c_lobby_abandon',
    trigger='用户在直播开始之前就关掉或离开了候场页（pagehide 事件触发）',
    value='分析流失原因：如果很多人在开播前 15 分钟放弃，说明候场时间设得太早了',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['duration_sec', '用户在候场页待了几秒才走的', '数字，例如：180（= 3分钟）'],
        ['minutes_until_start', '用户走的时候距离直播开始还有几分钟',
         '数字，例如：15（= 还有15分钟才开播就走了）'],
    ],
    code="trackGA4('c_lobby_abandon', {\n"
         "  webinar_id: '1',\n"
         "  duration_sec: 180,\n"
         "  minutes_until_start: 15\n"
         "});",
)

add_event_card(doc,
    event_name='c_add_to_calendar',
    trigger='用户点了「加到日历」按钮（支持 iCal 下载和 Google Calendar 两种方式）',
    value='衡量参与意愿：加了日历的人出席率通常很高，可以做为高意向受众',
    params=[
        ['method', '用户选择了哪种日历方式',
         '"ical" = 下载 .ics 文件（Apple/Outlook）\n"google" = 打开 Google Calendar 添加'],
        ['webinar_id', '直播场次编号', '"1"'],
    ],
    code="trackGA4('c_add_to_calendar', {\n"
         "  method: 'google',\n"
         "  webinar_id: '1'\n"
         "});",
)


doc.add_page_break()


# ── ④ 观看直播 ──
doc.add_heading('④ 观看直播', level=2)
doc.add_paragraph(
    '直播间是整个平台最核心的页面，追踪也最细致。'
    '这些数据可以帮你了解直播内容的效果。'
)

add_event_card(doc,
    event_name='c_video_heartbeat',
    trigger='视频播放中，每 60 秒自动记录一次；视频结束时也会再记一次（最终心跳）',
    value='像心跳一样持续记录「用户还在看吗」，可以精确算出每个人看了多久',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['current_time_sec', '视频播放到了第几秒', '数字，例如：1200（= 播到第 20 分钟）'],
        ['watch_duration_sec', '用户实际观看了几秒\n（扣除迟到加入的时间偏移）',
         '数字，例如：900\n（虽然视频在第 20 分钟，\n但用户是第 5 分钟才进来的，\n实际看了 15 分钟 = 900 秒）'],
    ],
    code="trackGA4('c_video_heartbeat', {\n"
         "  webinar_id: '1',\n"
         "  current_time_sec: 1200,\n"
         "  watch_duration_sec: 900\n"
         "});\n"
         "// 文件位置：src/hooks/usePlaybackTracking.ts",
)

add_event_card(doc,
    event_name='c_video_progress',
    trigger='视频播放到每 5% 里程碑时记一次（5%, 10%, 15%... 直到 100%），共 20 个里程碑，每个只触发一次',
    value='可以画出「观众留存曲线」：80% 的人看到了 30 分钟，但只有 40% 看完全场',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['percent', '视频播放到了百分之几', '数字：5, 10, 15, 20 ... 95, 100'],
    ],
    code="trackGA4('c_video_progress', {\n"
         "  webinar_id: '1',\n"
         "  percent: 50\n"
         "});\n"
         "// 文件位置：src/hooks/usePlaybackTracking.ts",
)

add_event_card(doc,
    event_name='c_chat_message',
    trigger='用户在聊天区发送了一条消息',
    value='互动参与度指标：发消息的人通常更投入，购买意愿也更高',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['video_time_sec', '用户发消息时视频正播到第几秒', '数字，例如：2400（= 第 40 分钟发的）'],
    ],
    code="trackGA4('c_chat_message', {\n"
         "  webinar_id: '1',\n"
         "  video_time_sec: 2400\n"
         "});\n"
         "// 文件位置：src/app/(public)/webinar/[id]/live/page.tsx",
)

add_event_card(doc,
    event_name='c_cta_view',
    trigger='购买推广弹窗（CTA）出现在用户屏幕上',
    value='多少人「看到了」购买提示？这个数字除以最终购买数就是弹窗的转化率',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['cta_id', '弹窗的唯一编号\n（后台设置的）', '字符串，例如："cta_offer_1"'],
        ['cta_type', '弹窗上的按钮文字\n（注意：虽然叫 cta_type，\n实际传的是按钮文字）',
         '字符串，最长 100 字\n例如："立即抢购 $599"'],
        ['video_time_sec', '弹窗出现时视频正播到第几秒', '数字，例如：1800（= 第 30 分钟弹出）'],
    ],
    code="trackGA4('c_cta_view', {\n"
         "  webinar_id: '1',\n"
         "  cta_id: 'cta_offer_1',\n"
         "  cta_type: '立即抢购 $599',\n"
         "  video_time_sec: 1800\n"
         "});",
)

add_event_card(doc,
    event_name='c_cta_dismiss',
    trigger='用户主动把购买弹窗关掉了',
    value='弹窗被关掉的比例：如果太高，可能弹窗出现的时机不对或内容不够吸引',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['cta_id', '被关掉的弹窗编号', '同 c_cta_view'],
        ['cta_type', '弹窗按钮文字', '同 c_cta_view'],
        ['video_time_sec', '用户关掉弹窗时视频播到第几秒', '数字'],
    ],
    code="trackGA4('c_cta_dismiss', {\n"
         "  webinar_id: '1',\n"
         "  cta_id: 'cta_offer_1',\n"
         "  cta_type: '立即抢购 $599',\n"
         "  video_time_sec: 1830\n"
         "});",
)

add_event_card(doc,
    event_name='begin_checkout（从直播间）',
    trigger='用户点了购买弹窗上的按钮，准备去结账',
    value='代表真实购买意图！记录了在视频第几秒点的，可以分析哪个时间点最容易成交',
    params=[
        ['currency', '货币', '固定值："USD"（美元）'],
        ['value', '产品金额', '数字：599'],
        ['items', '产品信息（GA4 电商标准格式）',
         '数组：[{\n  item_id: "webinar_1",\n  item_name: "课程名称",\n  price: 599,\n  quantity: 1\n}]'],
        ['cta_id', '用户点的是哪个弹窗', '字符串，例如："cta_offer_1"'],
        ['video_time_sec', '用户点购买时视频播到第几秒', '数字，例如：1842（= 30分42秒）'],
        ['source', '从哪个页面点的购买', '固定值："live"（直播间）'],
        ['--- 以下为自动附加的归因参数 ---', '', ''],
        ['gclid / utm_*', '广告来源信息', '同 sign_up'],
    ],
    code="trackGA4('begin_checkout', {\n"
         "  currency: 'USD',\n"
         "  value: 599,\n"
         "  items: [{ item_id: 'webinar_1', item_name: '课程名称', price: 599, quantity: 1 }],\n"
         "  cta_id: 'cta_offer_1',\n"
         "  video_time_sec: 1842,\n"
         "  source: 'live'\n"
         "});\n"
         "// 文件位置：src/app/(public)/webinar/[id]/live/page.tsx",
    is_conversion=True,
)

add_callout(doc,
    '亮点：视频追踪做得非常细 — 5% 的里程碑粒度加上 60 秒的心跳机制，'
    '足以支持精确的「观众留存曲线」分析。begin_checkout 记录了视频时间点，'
    '可以精确分析「直播的第几分钟最容易触发购买」。'
)


doc.add_page_break()


# ── ⑤ 结束页 ──
doc.add_heading('⑤ 直播结束 / 销售页', level=2)
doc.add_paragraph(
    '直播结束后，用户会看到一个销售页面，这是最后的成交机会。\n'
    '文件位置：src/app/(public)/webinar/[id]/end/page.tsx'
)

add_event_card(doc,
    event_name='c_webinar_complete',
    trigger='用户到达直播结束页面（代表用户看完了整场直播或看到了最后）',
    value='完播率 = 这个事件次数 ÷ c_enter_live 次数。是衡量直播内容质量的核心指标',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['watch_duration_sec', '用户总共看了多少秒的直播\n（从 sessionStorage 读取，\n可能为空）',
         '数字，例如：3600（= 看了 60 分钟）\n或 空（如果存储数据丢失）'],
        ['--- 以下为自动附加的归因参数 ---', '', ''],
        ['gclid / utm_*', '广告来源信息', '同 sign_up'],
    ],
    code="trackGA4('c_webinar_complete', {\n"
         "  webinar_id: '1',\n"
         "  watch_duration_sec: 3600\n"
         "  // + gclid, utm_* (自动附加)\n"
         "});",
    is_conversion=True,
)

add_event_card(doc,
    event_name='c_end_page_cta_click',
    trigger='用户在结束页面点了购买按钮',
    value='结束页的转化推动力：有些人是看完全场直播后在结束页才下定决心购买的',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['button_text', '用户点的按钮上写的文字', '字符串，例如："立即购买课程"'],
        ['--- 以下为自动附加的归因参数 ---', '', ''],
        ['gclid / utm_*', '广告来源信息', '同 sign_up'],
    ],
    code="trackGA4('c_end_page_cta_click', {\n"
         "  webinar_id: '1',\n"
         "  button_text: '立即购买课程'\n"
         "  // + gclid, utm_* (自动附加)\n"
         "});",
    is_conversion=True,
)

add_event_card(doc,
    event_name='begin_checkout（从结束页）',
    trigger='用户在结束页点了购买按钮（和 c_end_page_cta_click 同时触发）',
    value='与直播间的 begin_checkout 是同一个事件名，但 source 不同，方便区分来源',
    params=[
        ['currency', '货币', '"USD"'],
        ['value', '产品金额', '599'],
        ['items', '产品信息', '同直播间的 begin_checkout'],
        ['source', '从哪个页面点的', '固定值："end"（结束页）'],
        ['--- 以下为自动附加的归因参数 ---', '', ''],
        ['gclid / utm_*', '广告来源信息', '同 sign_up'],
    ],
    code="trackGA4('begin_checkout', {\n"
         "  currency: 'USD', value: 599,\n"
         "  items: [{ item_id: 'webinar_1', item_name: '课程名称', price: 599, quantity: 1 }],\n"
         "  source: 'end'\n"
         "});",
    is_conversion=True,
)

add_event_card(doc,
    event_name='c_share_click',
    trigger='用户点了社交媒体分享按钮（Facebook 或 Twitter）',
    value='口碑传播指标。分享功能使用率通常很低，但分享的用户是高满意度用户',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['platform', '用户分享到了哪个平台',
         '"facebook" = 分享到 Facebook\n"twitter" = 分享到 Twitter'],
    ],
    code="trackGA4('c_share_click', {\n"
         "  webinar_id: '1',\n"
         "  platform: 'facebook'\n"
         "});",
)


doc.add_page_break()


# ── ⑥ 付款完成 ──
doc.add_heading('⑥ 付款 / 购买完成', level=2)
doc.add_paragraph(
    '用户完成 Stripe 付款后，会看到一个确认页面。这里有最核心的转化事件。\n'
    '文件位置：src/app/(public)/checkout/[webinarId]/return/page.tsx'
)

add_event_card(doc,
    event_name='purchase',
    trigger='Stripe 确认付款成功，且确认页面完成加载。做了防重复处理（同一次付款只触发一次）',
    value='最最最重要的转化事件！Google Ads 用这个来计算 ROAS（花了多少广告费、赚了多少钱）',
    params=[
        ['transaction_id', 'Stripe 的交易编号\n（唯一标识这笔交易）',
         '字符串，例如："cs_live_abc123"\n如果 Stripe 没返回则用时间戳\n"session_1710288000000"'],
        ['value', '实际付款金额（美元）',
         '数字，例如：599\n来自 Stripe session 的 amountTotal ÷ 100\n如果 Stripe 没返回金额则默认 599'],
        ['currency', '货币代码', '"USD"（如果 Stripe 没返回则默认 USD）'],
        ['items', '购买的产品信息\n（GA4 电商标准格式）',
         '数组：[{\n  item_id: "webinar_1",\n  item_name: "课程名称",\n  price: 599,\n  quantity: 1\n}]'],
        ['--- 以下为自动附加的归因参数 ---', '', ''],
        ['gclid / utm_*', '广告来源信息', '同 sign_up'],
    ],
    code="trackGA4('purchase', {\n"
         "  transaction_id: 'cs_live_abc123',\n"
         "  value: 599,\n"
         "  currency: 'USD',\n"
         "  items: [{ item_id: 'webinar_1', item_name: '课程名称', price: 599, quantity: 1 }]\n"
         "  // + gclid, utm_* (自动附加)\n"
         "});",
    is_conversion=True,
)

add_event_card(doc,
    event_name='c_purchase_confirmation',
    trigger='和 purchase 同时触发（双重保险，备份用）',
    value='多记一次确保不漏，包含订单状态可以和后台数据库交叉验证',
    params=[
        ['webinar_id', '直播场次编号', '"1"'],
        ['transaction_id', 'Stripe 交易编号', '同 purchase 事件'],
        ['order_status', '订单状态（来自后台数据库）',
         '字符串，例如："completed" / "pending"\n如果查不到则为 "unknown"'],
    ],
    code="trackGA4('c_purchase_confirmation', {\n"
         "  webinar_id: '1',\n"
         "  transaction_id: 'cs_live_abc123',\n"
         "  order_status: 'completed'\n"
         "});",
)

add_callout(doc,
    '设计亮点：purchase 事件做了防重复（同一次付款只记一次），还有一个备份事件做双重保险。\n'
    '风险提醒：如果用户付完款但还没等页面加载完就关了浏览器，这笔购买就不会被 GA4 记录到。'
    '（第七章 P0-2 会详细说明解决方案）',
)


doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║       四、转化事件总表                                     ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('四、转化事件总表：Google Ads 最在意的 6 个动作', level=1)

doc.add_paragraph(
    '在所有 19 个追踪事件中，有 6 个被标记为「转化事件」。\n\n'
    '这些事件有两个特殊待遇：\n'
    '  1. 它们会自动带上广告来源信息（gclid + UTM），让 Google 知道这个转化是从哪个广告来的\n'
    '  2. 它们可以导入 Google Ads 作为优化目标，让机器学习去找更多「会做这个动作」的人'
)

add_table(doc,
    ['转化事件', '白话翻译', '在 Google Ads 中的角色', '建议转化价值'],
    [
        ['sign_up', '用户注册了',
         '「主要转化」\n让 Google 去找更多会注册的人',
         '建议设 $5~10\n（每个注册值多少钱？）'],
        ['c_enter_live', '用户进入了直播间',
         '「观察转化」\n辅助数据，帮 Google 理解漏斗',
         '建议设 $2'],
        ['begin_checkout', '用户点了购买按钮',
         '「观察转化」\n识别有购买意图的用户',
         '不设金额（仅观察）'],
        ['c_webinar_complete', '用户看完了整场直播',
         '「观察转化」\n衡量内容质量',
         '不设金额（仅观察）'],
        ['c_end_page_cta_click', '用户在结束页点了购买',
         '「观察转化」',
         '不设金额'],
        ['purchase', '用户付款成功了',
         '「主要转化」（最重要！）\n这是计算 ROAS 的核心',
         '动态金额（实际付了多少）'],
    ]
)

add_callout(doc,
    '建议在 Google Ads 中设置两个「主要转化」：sign_up（注册）和 purchase（购买）。\n'
    '其余 4 个设为「观察转化」（Secondary），不参与出价但可以看到数据。\n'
    '如果初期购买数据太少（每月 < 30 笔），可以先用 sign_up 作为主要优化目标。'
)


doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║       五、广告来源追踪                                     ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('五、广告来源追踪：我们是怎么记住「用户从哪来」的', level=1)

doc.add_paragraph(
    '当用户通过广告点进来，网址会带上标记参数。我们的网站会自动抓取并保存这些参数：'
)

add_table(doc,
    ['参数名', '是什么', '例子', '目前有没有抓'],
    [
        ['gclid', 'Google Ads 广告点击 ID\n（Google 自动加的）',
         '?gclid=EAIaIQ...', '有 ✓'],
        ['utm_source', '流量来源', '?utm_source=google', '有 ✓'],
        ['utm_medium', '流量媒介', '?utm_medium=cpc', '有 ✓'],
        ['utm_campaign', '广告活动名称', '?utm_campaign=spring_sale', '有 ✓'],
        ['utm_content', '广告素材标记', '?utm_content=video_ad_a', '有 ✓'],
        ['utm_term', '搜索关键字', '?utm_term=投资课程', '没有 ✗\n无法分析哪个关键字最赚钱'],
        ['wbraid', 'Google Ads iOS 追踪参数\n（Apple 隐私政策后的替代方案）',
         '?wbraid=CLaC...', '没有 ✗\niPhone 用户的广告效果无法追踪'],
        ['gbraid', 'Google Ads iOS App 追踪参数',
         '?gbraid=CJaD...', '没有 ✗\n同上'],
        ['fbclid', 'Facebook 广告点击 ID',
         '?fbclid=IwAR3...', '没有 ✗\n如果有投 Facebook 广告会需要'],
    ]
)

doc.add_paragraph(
    '保存方式：\n'
    '  • 第一份存在浏览器的临时存储（关浏览器就没了）— 用于当次访问\n'
    '  • 第二份存在 Cookie 里（保留 90 天）— 用于用户关掉浏览器后再回来\n\n'
    '这样即使用户今天看了着陆页，明天才注册，我们还是能知道他是从哪个广告来的。'
)


doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║       六、再营销受众                                       ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('六、再营销受众：可以圈出哪些人群再投广告', level=1)

doc.add_paragraph(
    '再营销（Retargeting）= 对「来过我们网站但没完成目标动作」的人再投广告。\n'
    '这比找完全陌生的新客人便宜 3~5 倍，转化率也高很多。\n\n'
    '基于目前的埋点事件，我们可以在 GA4 中自动圈出以下人群：'
)

add_table(doc,
    ['受众名称', '谁是这群人', '投什么广告给他们', '能不能建'],
    [
        ['看了没注册',
         '来了着陆页但没填注册表的人',
         '提醒广告：「限时免费直播，名额有限」',
         '可以 ✓'],
        ['注册没出席',
         '注册了但没来看直播的人',
         '重新邀请：「你错过了！下一场在...」',
         '可以 ✓'],
        ['出席没看完',
         '进了直播间但中途走了的人',
         '内容回顾：「精彩部分你可能错过了」',
         '可以 ✓'],
        ['看完没买',
         '看完全场直播但没付款的人',
         '限时优惠广告（最高价值受众！）',
         '可以 ✓'],
        ['加了购物车没买',
         '点了购买按钮但没完成付款的人',
         '购物车提醒：「你的课程还在等你」',
         '可以 ✓'],
        ['已购买客户',
         '已经付费的客户',
         '可以排除（不再花钱投给他们）\n或用来建 Lookalike 受众',
         '可以 ✓'],
        ['高参与观众',
         '看了 50% 以上直播的人',
         '类似受众扩展，找更多这样的人',
         '可以 ✓'],
    ]
)

add_callout(doc,
    '行动建议：请在 GA4 后台 → 受众 中建立以上 7 个受众，并连接到 Google Ads 的受众列表。'
    '再营销广告通常是 ROAS 最高的广告类型。'
)


doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║       七、健康检查                                         ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('七、健康检查：目前做得好的 & 还缺什么', level=1)

doc.add_heading('整体评分：7.5 / 10', level=2)
doc.add_paragraph(
    '核心追踪已经相当完善，基本能支持广告投放和数据分析。'
    '但有 3 个关键缺失如果不补，会导致数据不准、广告优化效果打折。'
)

doc.add_heading('做得好的部分', level=2)
good_items = [
    ('漏斗完整',
     '从着陆页到付款，每一步都有追踪，没有明显盲区'),
    ('广告来源自动跟踪',
     '6 个转化事件都会自动带上 gclid/UTM 信息，不用担心「不知道购买从哪个广告来」'),
    ('购买事件很完整',
     '记录了实际金额、交易编号、产品信息，GA4 可以直接算收入'),
    ('防丢设计',
     '广告来源存了两份（临时 + Cookie 90 天），购买事件有双重保险'),
    ('视频追踪很细',
     '每 5% 记一次进度 + 每 60 秒心跳，可以画出精确的观众留存曲线'),
    ('候场页分析完善',
     '不仅记录了有没有来，还记录了等了多久、在直播前几分钟放弃的'),
    ('代码品质好',
     '用了类型安全的代码（TypeScript），减少了写错事件名的可能性'),
]
for title, desc in good_items:
    p = doc.add_paragraph()
    run = p.add_run(f'✓ {title}：')
    run.bold = True
    run.font.color.rgb = GREEN
    p.add_run(desc)


doc.add_page_break()


# ── P0 ──
doc.add_heading('P0 — 必须马上补（没有会亏钱）', level=2)

doc.add_paragraph(
    '以下 3 个问题会直接影响广告花费的效率，建议尽快修复：'
)

# P0-1
add_bold_paragraph(doc, '问题一：注册和购买事件没有传用户资料给 Google（Enhanced Conversions 缺失）')
doc.add_paragraph(
    '现在的情况：\n'
    '用户注册时，我们收集了 email 和手机号。但这些资料只存到了自己的数据库，'
    '没有传给 Google。\n\n'
    '这为什么是个问题？\n'
    'Google Ads 需要用户的 email/手机来做「增强型转化」匹配。简单说就是：\n'
    '  • 用户在手机上看了广告，但在电脑上注册了 → Google 靠 email 把这两件事对起来\n'
    '  • Apple iPhone 限制了广告追踪 → Google 靠 email 做替代方案\n\n'
    '不修的后果：\n'
    '  • 转化数据可能少报 30~50%（明明有人从广告来注册了，但 Google 不知道）\n'
    '  • Google Ads 的「智能出价」效果变差（因为数据不够准）\n'
    '  • 你以为某个广告不行其实它很行，浪费了预算\n\n'
    '修复方法：\n'
    '在 sign_up 和 purchase 事件中，加上用户的 email（Google 会自动加密处理，不用担心隐私）。'
)

# P0-2
add_bold_paragraph(doc, '问题二：购买只靠浏览器记录，用户关掉浏览器就漏了（缺服务端追踪）')
doc.add_paragraph(
    '现在的情况：\n'
    '用户在 Stripe 付款成功后，需要回到我们的确认页面，页面加载时才会触发 purchase 事件。\n\n'
    '这为什么是个问题？\n'
    '如果用户付完钱后：\n'
    '  • 手机断网了\n'
    '  • 浏览器闪退了\n'
    '  • 不耐烦直接关掉了\n'
    '  • 装了广告拦截器（AdBlock），GTM 被阻止了\n'
    '→ 这笔购买就不会被 GA4 记录到！\n\n'
    '不修的后果：\n'
    '  • GA4 里的购买次数比实际少\n'
    '  • ROAS 看起来比实际低\n'
    '  • 可能错误地暂停了其实赚钱的广告\n\n'
    '修复方法：\n'
    '在服务器收到 Stripe 的付款通知（webhook）时，直接从服务器发一笔 purchase 事件给 GA4，'
    '不依赖用户的浏览器。这叫「Measurement Protocol」，是 GA4 提供的服务端 API。'
)

# P0-3
add_bold_paragraph(doc, '问题三：iPhone 用户的广告来源追踪不到（缺 wbraid/gbraid）')
doc.add_paragraph(
    '现在的情况：\n'
    '我们的网站会抓取网址中的 gclid（Google 广告点击 ID），但没有抓 wbraid 和 gbraid。\n\n'
    '这为什么是个问题？\n'
    'Apple 从 iOS 14 开始限制了广告追踪。Google 的应对方案是在 iPhone 用户的网址上'
    '用 wbraid/gbraid 代替 gclid。如果我们不抓这两个参数：\n'
    '  • 所有 iPhone 用户的转化都无法归因到具体广告\n'
    '  • 在北美市场，iPhone 用户占比约 50-60%\n'
    '  • 等于一半用户的广告效果完全看不到\n\n'
    '修复方法：\n'
    '在广告来源捕获组件（GclidPreserver）中加上 wbraid 和 gbraid 两个参数，改动非常小。'
)


doc.add_page_break()


# ── P1 ──
doc.add_heading('P1 — 建议尽快补（能明显提升效果）', level=2)

add_table(doc,
    ['#', '缺什么', '影响', '怎么补'],
    [
        ['P1-1',
         '注册后没有设 GA4 的 user_id',
         '同一个人在手机和电脑上看到的，\nGA4 会算成两个人',
         '用户注册后，把 email 的加密值\n设为 GA4 的 user_id'],
        ['P1-2',
         '不知道注册弹窗有没有真的弹出来\n（缺 c_registration_modal_open 事件）',
         '无法计算「点了报名按钮」到\n「弹窗真的打开了」之间的流失',
         '弹窗打开时触发一个事件'],
        ['P1-3',
         '不知道用户有没有真的到结账页\n（缺 c_checkout_page_view 事件）',
         '知道有人点了购买，但不知道\nStripe 结账页有没有成功加载',
         '结账页加载时触发一个事件'],
        ['P1-4',
         '不知道视频有没有真的开始播\n（缺 c_video_play_start 事件）',
         '无法区分「进了直播间」和\n「视频真的在播了」',
         '视频首次播放时触发一个事件'],
        ['P1-5',
         '没有抓 utm_term\n（搜索关键字参数）',
         '无法分析哪个关键字\n带来的注册和购买最多',
         '在 GclidPreserver 中\n加上 utm_term'],
        ['P1-6',
         '不知道用户有没有开声音\n（缺 c_unmute_click 事件）',
         '无法分辨「真的在看」和\n「只是挂在后台」',
         '用户取消静音时\n触发一个事件'],
    ]
)


# ── P2 ──
doc.add_heading('P2 — 锦上添花（有更好，没有不影响核心功能）', level=2)

add_table(doc,
    ['#', '可以加的功能', '用途'],
    [
        ['P2-1', '着陆页独立浏览事件', '更精确地追踪着陆页访问（不依赖 GTM 自动追踪）'],
        ['P2-2', '重播点击事件', '如果有重播功能，可以追踪有多少人想重看'],
        ['P2-3', '错过场次 / 重新预约', '追踪常青直播系统中用户错过和重新预约的行为'],
        ['P2-4', '侧边栏标签切换', '了解用户在直播间喜欢看哪些内容（聊天？信息？优惠？）'],
        ['P2-5', '标签页切换', '用户是否切到了其他标签页（衡量注意力）'],
        ['P2-6', '注册表单填写错误', '如果很多人在手机号格式上犯错，可以优化提示语'],
        ['P2-7', '抓取 fbclid', '如果未来投 Facebook 广告，需要这个来追踪 FB 的转化'],
        ['P2-8', 'Google Ads 离线转化导入', '把服务器上的购买数据回传给 Google Ads'],
        ['P2-9', '结账放弃追踪', '用户到了付款页但没完成付款'],
        ['P2-10', '结束页滚动深度', '用户在最终销售页看了多少内容'],
    ]
)


doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║       八、给技术同事的清单                                  ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('八、GTM / GA4 配置端要做的事（给技术同事的清单）', level=1)

doc.add_paragraph(
    '以下是在 GTM 和 GA4 后台需要完成的配置工作。'
    '这些不涉及网页代码修改，而是在 Google 的管理界面里操作：'
)

config_items = [
    ('GTM 容器',
     '确认 GTM 容器中已正确配置 GA4 Configuration Tag，使用正确的 Measurement ID（G-XXXXX）'),
    ('转化标记',
     '在 GA4 管理界面 → 转化事件 中，将 sign_up 和 purchase 标记为「转化事件」'),
    ('Google Ads 转化',
     '在 Google Ads 中导入 GA4 的 sign_up（注册转化）和 purchase（购买转化），设置对应的转化价值'),
    ('Enhanced Conversions',
     '在 GTM 中配置 Enhanced Conversions 变量，读取 dataLayer 中的 user_data（等开发端修复 P0-1 后）'),
    ('再营销标签',
     '配置 Google Ads Remarketing Tag，使用第六章定义的 7 个受众群体'),
    ('自定义定义',
     '在 GA4 管理界面 → 自定义定义 中，注册所有 c_ 前缀的事件参数为自定义维度/指标'),
    ('Consent Mode',
     '如果需要遵守隐私法规（GDPR/CCPA），配置 Google Consent Mode'),
]

for i, (title, desc) in enumerate(config_items, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title}：')
    run.bold = True
    p.add_run(desc)


doc.add_heading('事件命名规则说明', level=2)
doc.add_paragraph(
    '在 GA4 后台，你会看到两类事件名：\n\n'
    '无前缀的事件（sign_up, begin_checkout, purchase）\n'
    '  → 这些是 Google 官方推荐的标准事件名，GA4 会自动识别，不用额外配置\n\n'
    'c_ 开头的事件（c_scroll_depth, c_enter_live, c_video_progress 等）\n'
    '  → 这些是我们自定义的事件（c = custom），需要在 GA4 管理界面中\n'
    '    手动注册为「自定义维度」才能在报表中使用\n\n'
    '所有事件名一览（方便在 GA4 后台核对）：'
)

add_table(doc,
    ['事件名', '类型', '是不是转化', '在哪个页面触发'],
    [
        ['c_scroll_depth', '自定义', '否', '着陆页'],
        ['c_signup_button_click', '自定义', '否', '着陆页'],
        ['sign_up', 'Google 标准', '是 ✓', '注册弹窗'],
        ['c_lobby_entered', '自定义', '否', '候场等待页'],
        ['c_enter_live', '自定义', '是 ✓', '候场→直播间'],
        ['c_lobby_duration', '自定义', '否', '候场等待页'],
        ['c_lobby_abandon', '自定义', '否', '候场等待页'],
        ['c_add_to_calendar', '自定义', '否', '候场等待页'],
        ['c_video_heartbeat', '自定义', '否', '直播间'],
        ['c_video_progress', '自定义', '否', '直播间'],
        ['c_chat_message', '自定义', '否', '直播间'],
        ['c_cta_view', '自定义', '否', '直播间'],
        ['c_cta_dismiss', '自定义', '否', '直播间'],
        ['begin_checkout', 'Google 标准', '是 ✓', '直播间 / 结束页'],
        ['c_webinar_complete', '自定义', '是 ✓', '结束页'],
        ['c_end_page_cta_click', '自定义', '是 ✓', '结束页'],
        ['c_share_click', '自定义', '否', '结束页'],
        ['purchase', 'Google 标准', '是 ✓', '付款确认页'],
        ['c_purchase_confirmation', '自定义', '否', '付款确认页'],
    ]
)


doc.add_page_break()


# ╔══════════════════════════════════════════════════════════╗
# ║       附录：完整代码文件索引                                ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading('附录：完整代码文件索引（给工程师看的）', level=1)

doc.add_paragraph(
    '如果工程师需要修改追踪代码，以下是所有相关文件的位置：'
)

add_table(doc,
    ['文件路径', '这个文件做什么'],
    [
        ['src/lib/analytics.ts', '追踪核心：trackGA4() 函数在这里\n所有事件参数的类型定义也在这里'],
        ['src/components/analytics/GclidPreserver.tsx', '广告来源捕获：\n抓 gclid/UTM 存到 Cookie'],
        ['src/app/layout.tsx', 'GTM 代码加载的地方'],
        ['src/types/global.d.ts', 'TypeScript 类型声明\n（window.dataLayer 的定义）'],
        ['src/app/(public)/page.tsx', '着陆页的追踪代码\n（scroll_depth, signup_button_click）'],
        ['src/components/registration/useRegistrationForm.ts', '注册事件（sign_up）在这里触发'],
        ['src/app/(public)/webinar/[id]/lobby/page.tsx', '候场页所有追踪\n（entered, enter_live, abandon 等）'],
        ['src/app/(public)/webinar/[id]/live/page.tsx', '直播间追踪\n（chat, cta_view, begin_checkout）'],
        ['src/hooks/usePlaybackTracking.ts', '视频播放追踪 Hook\n（heartbeat, progress）'],
        ['src/app/(public)/webinar/[id]/end/page.tsx', '结束页追踪\n（webinar_complete, share_click）'],
        ['src/app/(public)/checkout/[webinarId]/return/page.tsx', '付款确认追踪\n（purchase, purchase_confirmation）'],
    ]
)


# ── Save ──
output_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'docs', 'webinar-mvp-tracking-audit-report.docx'
)
doc.save(output_path)
print(f"Report saved to: {output_path}")
